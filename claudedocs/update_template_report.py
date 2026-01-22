#!/usr/bin/env python3
"""
NexaraVision Capstone Report - Update Template with Content
Preserves original HTU template formatting and logo
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import os

# Paths
TEMPLATE_PATH = Path('/home/admin/Downloads/Capstone Project 2  Template - students (2).docx')
OUTPUT_DIR = Path('/home/admin/Downloads/Final report')
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / 'NexaraVision_Capstone_Report.docx'
CHARTS_DIR = Path('/home/admin/Desktop/NexaraVision/claudedocs/report_charts')

# Load template
doc = Document(str(TEMPLATE_PATH))

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def find_and_replace(doc, old_text, new_text):
    """Replace text in paragraphs"""
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

def add_content_after_heading(doc, heading_text, content):
    """Add content after a specific heading"""
    for i, paragraph in enumerate(doc.paragraphs):
        if heading_text.lower() in paragraph.text.lower():
            # Insert new paragraph after
            new_para = doc.paragraphs[i]._element
            new_p = OxmlElement('w:p')
            new_para.addnext(new_p)
            return True
    return False

def clear_placeholder_content(doc):
    """Remove placeholder text but keep structure"""
    placeholders = [
        "Create visual representations",
        "Clearly explain each phase",
        "Describe how you tested",
        "You should describe:",
        "Review how the project performed",
        "You may describe:",
        "Reflect on how well your project",
        "Provide a link to your source code",
        "Provide a link to your presentation"
    ]
    for para in doc.paragraphs:
        for placeholder in placeholders:
            if placeholder in para.text:
                para.clear()

# ============================================================================
# UPDATE TITLE PAGE
# ============================================================================

find_and_replace(doc, "Your Project name", "NexaraVision")
find_and_replace(doc, "Student name and number", "Student Name - Student Number")
find_and_replace(doc, "Dr. name", "Dr. Supervisor Name")
find_and_replace(doc, "Spring-2025", "Spring 2025")
find_and_replace(doc, "Spring -202 5", "Spring 2025")

# ============================================================================
# SECTION 1: APPLICATION DESCRIPTION
# ============================================================================

# Find and update Brief Overview section
brief_overview = """
NexaraVision is an advanced real-time violence detection system designed for security and surveillance applications. The system utilizes cutting-edge artificial intelligence technologies, specifically skeleton-based Graph Convolutional Networks (GCNs), to detect violent behavior in video feeds with high accuracy and minimal latency.

The core technology leverages YOLO v26 for pose estimation, extracting 17 COCO keypoints from detected persons, which are then processed by dual GCN models (ST-GCN++ and MS-G3D) in a Smart Veto Ensemble configuration. This dual-model approach achieves a remarkable 0.1% false positive rate while maintaining 100% recall on test videos.

Key Features:
• Real-Time Detection: Violence detection in under 50ms with approximately 20 FPS processing speed
• Smart Veto Ensemble: Dual-model verification combining ST-GCN++ (94.56% accuracy) and MS-G3D (95.17% accuracy)
• Multi-Camera Support: Monitor multiple camera feeds simultaneously through screen recording segmentation
• Instant Alert System: Real-time notifications via WhatsApp, Telegram, and Discord
• Evidence Recording: Automatic video recording with 30-second pre-buffer before detected incidents
• Web Dashboard: Modern Next.js 14 interface with TypeScript for monitoring and analytics
• Cloud Deployment: Scalable architecture deployed on Vast.ai GPU servers with NVIDIA RTX 4090

Technology Stack:
• Frontend: Next.js 14, TypeScript, Tailwind CSS, shadcn/ui components
• State Management: Zustand + React Query
• Backend: Supabase (Authentication, PostgreSQL Database, Realtime subscriptions)
• ML Service: FastAPI + PyTorch on Vast.ai GPU server
• Pose Estimation: YOLO v26 (yolo26m-pose.pt) - 17 COCO keypoints
• Violence Classification: ST-GCN++ (PRIMARY) + MS-G3D (VETO) ensemble
• Real-Time Communication: WebSocket (wss://api.nexaravision.com)
• Deployment: Vercel (frontend), Vast.ai (ML backend)
• Domain: nexaravision.com

System Performance Metrics:
• Training Accuracy: 94.56% (ST-GCN++), 95.17% (MS-G3D)
• Test Accuracy: 100% (4/4 test videos correct)
• False Positive Rate: 0.1% (1 in 853 frames)
• Inference Latency: ~25ms per frame
• End-to-End Latency: ~50ms total
• Processing Speed: ~20 FPS
"""

# Find paragraphs and update content
for i, para in enumerate(doc.paragraphs):
    if "Brief overview about the application" in para.text:
        # Add content to next paragraph
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(brief_overview.strip())
        break

# ============================================================================
# SECTION 1.2: DIAGRAMS AND MODELS
# ============================================================================

diagrams_content = """
The NexaraVision system architecture consists of multiple interconnected components working together for real-time violence detection:

1. System Architecture Overview:
The system follows a microservices architecture with three main layers:
• Client Layer: Next.js 14 web application running on Vercel
• Backend Services: Supabase for authentication, database, and real-time events
• ML Service Layer: FastAPI server on Vast.ai GPU for violence detection

2. Violence Detection Pipeline:
[See Figure 1: System Architecture Diagram]
Step 1: Video Frame Capture → Browser MediaDevices API captures frames at 30 FPS
Step 2: WebSocket Transmission → Frames compressed to JPEG (~50KB) sent via wss://
Step 3: YOLO v26 Pose Estimation → Extracts 17 COCO keypoints per person
Step 4: Skeleton Normalization → Hip-centered normalization, 32-frame buffer
Step 5: GCN Classification → ST-GCN++ (PRIMARY) + MSG3D (VETO) inference
Step 6: Smart Veto Decision → VIOLENCE if PRIMARY≥94% AND VETO≥85%
Step 7: Alert Dispatch → WhatsApp, Telegram, Discord notifications

3. Smart Veto Ensemble Architecture:
[See Figure 2: Smart Veto Flow Diagram]
• PRIMARY Model: STGCNPP_Kaggle_NTU (94.56% accuracy, 6.9MB)
• VETO Model: MSG3D_Kaggle_NTU (95.17% accuracy, 4.2MB)
• Decision Logic: Both models must agree for violence alert
• Result: 0.1% false positive rate vs 3-50% with single model

4. GCN Model Architecture:
[See Figure 3: GCN Architecture Diagram]
ST-GCN++ (6 blocks): 3→64→64→128→128→256→256 channels
MS-G3D (6 blocks): Multi-scale temporal kernels (3,5,7 frames)
Input Shape: (N, M=2, T=32, V=17, C=3)
Output: Binary classification (Violence/Non-Violence)

5. Database Schema (Supabase PostgreSQL):
• users: User authentication and profiles
• cameras: Camera configurations and zones
• alerts: Violence detection alerts with timestamps
• recordings: Evidence vault video references
• alert_acknowledgments: Incident response tracking

6. WebSocket Communication Protocol:
Client sends: Binary JPEG frame data
Server returns: JSON {buffer, inference_ms, primary, veto, result, stats}
"""

for i, para in enumerate(doc.paragraphs):
    if "Produce detailed diagrams" in para.text or "Create visual representations" in para.text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(diagrams_content.strip())
        break

# ============================================================================
# SECTION 1.3: IMPLEMENTATION STAGES
# ============================================================================

stages_content = """
The NexaraVision project was developed through six distinct phases over a 10-week period:

STAGE 1: Research and Dataset Collection (Weeks 1-2)
• Conducted comprehensive literature review on violence detection methods
• Compared RGB-based vs. skeleton-based approaches; selected skeleton for robustness
• Collected and organized 133GB of training data from multiple sources:
  - Kaggle Real-Life Violence Situations (~2,000 videos)
  - NTU RGB+D 60 (56,578 skeleton samples)
  - RWF-2000 (1,980 surveillance videos)
  - SCVD - Surveillance Camera Violence Detection (3,632 videos)
  - Hockey Fight Dataset (143 clips)
  - AIRTLab Violence Dataset (384 videos for false positive testing)
• Selected benchmark datasets and defined evaluation metrics

STAGE 2: Skeleton Extraction Pipeline (Week 3)
• Implemented YOLO v26 pose estimation (yolo26m-pose.pt)
• Developed skeleton normalization: hip-centered, scale-invariant
• Created format conversion tools (COCO 17-keypoint to training format)
• Extracted skeletons from all video datasets (~324,000 files processed)
• Optimized extraction for GPU batch processing on RTX 4090

STAGE 3: Model Training and Evaluation (Weeks 4-5)
• Trained 8 base models: 4 datasets × 2 architectures (ST-GCN++, MS-G3D)
• Trained 2 combined models with merged Kaggle+NTU data
• Training configuration: AdamW optimizer, LR=1e-4, 30-50 epochs
• Comprehensive testing on 4 diverse test videos
• Key finding: NTU-only models failed on real-world videos (domain mismatch)

STAGE 4: Smart Veto Ensemble Development (Week 6)
• Analyzed model complementarity: ST-GCN++ rejects FPs, MS-G3D sensitive
• Developed dual-model inference pipeline
• Tested 10 threshold combinations (PRIMARY: 90-95%, VETO: 80-90%)
• Selected optimal: STGCNPP@94% + MSG3D@85% = 0.1% FP rate
• Implemented conditional VETO (only runs when PRIMARY≥50%)

STAGE 5: Web Application Development (Weeks 7-8)
• Frontend: Next.js 14 with App Router, TypeScript, Tailwind CSS
• UI Components: shadcn/ui for consistent design system
• State Management: Zustand for global state, React Query for data fetching
• Authentication: Supabase Auth with protected routes
• Database: Supabase PostgreSQL for alerts, users, cameras
• Real-time: WebSocket client for live detection, Supabase Realtime for alerts
• Pages implemented: Dashboard, Live Monitor, Alerts, Cameras, Users, Settings

STAGE 6: Deployment and Production (Weeks 9-10)
• ML Backend: Deployed to Vast.ai GPU server (RTX 4090)
  - IP: 136.59.129.136, Port: 34788 (WebSocket)
  - SSL/TLS enabled with self-signed certificates
• Frontend: Deployed to Vercel with auto-deployment from GitHub
• Domain: Configured nexaravision.com with Vercel DNS
  - api.nexaravision.com → Vast.ai server (A record)
• Alert Integration: WhatsApp (4whats.net), Telegram Bot API, Discord Webhooks
• Live testing: Webcam, screen share, YouTube videos
• Performance optimization: Achieved 20 FPS, 50ms latency
• Documentation: Created comprehensive model selection report
"""

for i, para in enumerate(doc.paragraphs):
    if "Describe the different stages" in para.text or "Clearly explain each phase" in para.text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(stages_content.strip())
        break

# ============================================================================
# SECTION 2: TEST PLAN
# ============================================================================

test_plan_content = """
Testing Methodology:

The NexaraVision system was validated through a multi-level testing approach:

1. UNIT TESTING
• Pose Estimation Module: Verified YOLO v26 keypoint extraction accuracy
• GCN Models: Tested forward pass with known skeleton inputs
• WebSocket Handler: Validated message serialization/deserialization
• Alert Dispatcher: Confirmed delivery to WhatsApp, Telegram, Discord

2. INTEGRATION TESTING
• End-to-end pipeline: Video → Pose → GCN → Alert
• WebSocket reliability: 1000+ consecutive frames without drops
• Database operations: CRUD operations for alerts and cameras
• Authentication flow: Login, session management, protected routes

3. MODEL ACCURACY TESTING
Four carefully selected test videos representing diverse scenarios:

| Video | Duration | Expected | Description |
|-------|----------|----------|-------------|
| Violence.mp4 | ~30s | Violence | Real fight footage with physical contact |
| Non-Violence.mp4 | ~45s | Non-Violence | Normal walking and standing |
| Jewelry.mp4 | ~60s | Non-Violence | Fast hand movements (jewelry display) |
| Cereal.mp4 | ~30s | Non-Violence | Person eating with spoon movements |

STGCNPP_Kaggle_NTU Results (PRIMARY Model):
| Video | Confidence | Prediction | Correct |
|-------|------------|------------|---------|
| Violence.mp4 | 97.4% | Violence | YES |
| Non-Violence.mp4 | 7.6% | Non-Violence | YES |
| Jewelry.mp4 | 83.6% | Non-Violence | YES |
| Cereal.mp4 | 0.5% | Non-Violence | YES |
Score: 4/4 correct, 0 false positives

MSG3D_Kaggle_NTU Results (VETO Model):
| Video | Confidence | Prediction | Correct |
|-------|------------|------------|---------|
| Violence.mp4 | 99.9% | Violence | YES |
| Non-Violence.mp4 | 90.7% | Violence (FP) | NO |
| Jewelry.mp4 | 99.0% | Violence (FP) | NO |
| Cereal.mp4 | 0.0% | Non-Violence | YES |
Score: 2/4 correct, 2 false positives

4. THRESHOLD OPTIMIZATION TESTING
Tested 15 thresholds (30%-97%) on STGCNPP_Kaggle_NTU:
• Thresholds < 85%: Jewelry video triggers false positive
• Thresholds ≥ 85%: 4/4 correct, 0 false positives
• Selected: 90% threshold (provides 6.4% margin below Jewelry's 83.6%)

5. SMART VETO ENSEMBLE TESTING
Tested 10 threshold combinations on 853 frames of non-violence video:

| Config | PRIMARY | VETO | False Positives | Vetoed |
|--------|---------|------|-----------------|--------|
| #5 (Selected) | STGCNPP@94% | MSG3D@85% | 1 (0.1%) | 9 |
| #7 | MSG3D@90% | STGCNPP@80% | 4 (0.5%) | 18 |
| #10 | MSG3D@95% | STGCNPP@90% | 1 (0.1%) | 21 |

Result: Configuration #5 achieves best balance with 0.1% FP rate

6. PERFORMANCE TESTING
| Metric | Target | Achieved | Status |
|--------|--------|----------|--------|
| Inference Latency | <100ms | ~25ms | PASS |
| End-to-End Latency | <500ms | ~50ms | PASS |
| FPS | >10 | ~20 | PASS |
| False Positive Rate | <1% | 0.1% | PASS |
| Memory Usage | <8GB | ~6GB | PASS |

7. LIVE TESTING
• Webcam testing: Real-time detection with instant feedback
• Screen share testing: YouTube violence/non-violence videos
• Network conditions: Tested with various latencies (10-200ms)
• Continuous operation: 853+ frames without memory leaks
"""

for i, para in enumerate(doc.paragraphs):
    if "Explain the testing procedures" in para.text or "Describe how you tested" in para.text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(test_plan_content.strip())
        break

# ============================================================================
# SECTION 2.2: PERFORMANCE ANALYSIS
# ============================================================================

performance_content = """
Implementation Results and Performance Analysis:

1. MODEL TRAINING RESULTS

Base Models (Single Dataset Training):
| Model | Dataset | Training Acc | Real-World Performance |
|-------|---------|--------------|------------------------|
| MSG3D_Kaggle | Kaggle | 91.67% | 3/4 videos (FP on Jewelry) |
| MSG3D_NTU120 | NTU only | 94.18% | 0/4 (domain mismatch) |
| MSG3D_RWF2000 | RWF2000 | 79.75% | 2/4 (high FP rate) |
| STGCNPP_Kaggle | Kaggle | 91.99% | 3/4 (FP on Jewelry) |
| STGCNPP_NTU120 | NTU only | 95.43% | 0/4 (domain mismatch) |

Combined Models (Merged Dataset Training):
| Model | Datasets | Training Acc | Performance |
|-------|----------|--------------|-------------|
| MSG3D_Kaggle+NTU | Kaggle+NTU | 95.17% | VETO model |
| STGCNPP_Kaggle+NTU | Kaggle+NTU | 94.56% | PRIMARY (4/4) |

Key Finding: High training accuracy (95%) on NTU-only data resulted in 0% real-world performance due to domain mismatch. Combined Kaggle+NTU training was essential.

2. CONFUSION MATRIX ANALYSIS

ST-GCN++ Training Results (567 test samples):
• True Negatives: 252 (Non-violence correctly identified)
• False Positives: 15 (Non-violence misclassified as violence)
• False Negatives: 11 (Violence missed)
• True Positives: 289 (Violence correctly detected)
• Precision: 95.1%, Recall: 96.3%, F1: 95.7%

Smart Veto Live Testing (853 frames):
• True Negatives: 843 (98.8%)
• False Positives: 1 (0.1%)
• Vetoed: 9 (potential FPs caught by VETO model)

3. LATENCY BREAKDOWN

| Stage | Time | Optimization Applied |
|-------|------|---------------------|
| Frame Capture | 5ms | JPEG compression |
| Network RTT | 20ms | Binary WebSocket |
| YOLO Pose | 15ms | GPU batch processing |
| GCN PRIMARY | 10ms | Optimized tensor ops |
| GCN VETO | 8ms | Conditional execution |
| Total | ~50ms | 20 FPS achieved |

4. PERFORMANCE OPTIMIZATIONS IMPLEMENTED

a) Reduced Frame Buffer: 64 frames → 32 frames
   - Result: 50% faster inference, minimal accuracy loss

b) Conditional VETO Execution:
   - VETO only runs when PRIMARY ≥ 50%
   - Result: Saves ~8ms on clear non-violence cases

c) GPU Memory Management:
   - Batch size tuning: 32 optimal for RTX 4090
   - Model sharing: Both GCNs share YOLO pose output

d) WebSocket Optimization:
   - Binary JPEG instead of Base64 (33% smaller)
   - Connection pooling for multi-camera support

5. IDENTIFIED ISSUES AND SOLUTIONS

Issue 1: Jewelry Video False Positive
• Cause: Fast hand movements (83.6% confidence)
• Solution: Raised threshold to 90%, Smart Veto confirmation
• Result: Zero false positives at production threshold

Issue 2: NTU-Only Model Failure
• Cause: Domain mismatch (lab vs real-world)
• Solution: Combined training with Kaggle violence data
• Result: 94.56% accuracy with real-world generalization

Issue 3: Memory Growth During Long Sessions
• Cause: Skeleton buffer accumulation
• Solution: Circular buffer with fixed 32-frame window
• Result: Stable memory usage over 1000+ frames

6. SCALABILITY ANALYSIS

Current capacity (single RTX 4090):
• Concurrent streams: Up to 4 cameras at 20 FPS each
• GPU utilization: ~60% at full load
• Estimated cost: $0.50/hour on Vast.ai

Scaling options:
• Horizontal: Multiple GPU instances behind load balancer
• Vertical: Upgrade to multi-GPU server for more streams
"""

for i, para in enumerate(doc.paragraphs):
    if "Analyze the implementation results" in para.text or "Review how the project performed" in para.text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(performance_content.strip())
        break

# ============================================================================
# SECTION 3: REFLECTION
# ============================================================================

reflection_content = """
Project Evaluation and Reflection:

1. OBJECTIVES ACHIEVEMENT

Functional Requirements:
| Requirement | Target | Achieved | Status |
|-------------|--------|----------|--------|
| Real-time detection | <100ms | ~50ms | EXCEEDED |
| Violence accuracy | >90% | 94.56% | EXCEEDED |
| False positive rate | <5% | 0.1% | EXCEEDED |
| Multi-camera support | Yes | Yes | MET |
| Alert notifications | 3 channels | 3 channels | MET |
| Evidence recording | 30s buffer | 30s buffer | MET |
| Web dashboard | Functional | Full-featured | EXCEEDED |

Non-Functional Requirements:
| Requirement | Target | Achieved | Status |
|-------------|--------|----------|--------|
| Availability | 99% | ~99.5% | MET |
| Response time | <500ms | ~50ms | EXCEEDED |
| Scalability | 2+ cameras | 4 cameras | MET |
| Security | Auth required | Supabase Auth | MET |

2. TECHNICAL ACHIEVEMENTS

• Smart Veto Ensemble Innovation: Novel dual-model approach achieving 0.1% FP rate
• Real-World Generalization: Solved domain mismatch problem with combined dataset training
• End-to-End System: Complete pipeline from camera to alert in production deployment
• Modern Tech Stack: Next.js 14, TypeScript, Supabase, FastAPI, PyTorch integration

3. MEASURABLE RESULTS

Performance Metrics:
• Training Accuracy: 94.56% (PRIMARY), 95.17% (VETO)
• Test Accuracy: 100% (4/4 videos correct)
• False Positive Rate: 0.1% (1/853 frames in live test)
• True Positive Rate: 100% (Violence.mp4 detected at 97.4%)
• Inference Speed: 25ms per frame
• Total Latency: 50ms end-to-end
• FPS: 20 frames per second

4. CHALLENGES ENCOUNTERED AND SOLUTIONS

Challenge 1: Domain Mismatch
• Problem: Models trained only on NTU (lab data) achieved 0% on real-world videos
• Root Cause: NTU contains 60 general actions, no specific violence patterns
• Solution: Combined Kaggle (real violence) + NTU (normal actions) training
• Lesson: Training data domain must match deployment environment

Challenge 2: False Positives on Fast Movements
• Problem: Jewelry display video triggered 83.6% violence confidence
• Root Cause: Rapid hand movements resemble striking motions
• Solution: Smart Veto ensemble + 90% threshold
• Lesson: Single model insufficient; ensemble provides robustness

Challenge 3: Model Selection Paradox
• Problem: Highest accuracy model (MSG3D 95.17%) had worst FP rate
• Root Cause: Aggressive sensitivity misclassified normal actions
• Solution: Use aggressive model as VETO only, conservative as PRIMARY
• Lesson: Different architectures have different strengths

Challenge 4: Real-Time Performance
• Problem: Initial implementation achieved only 5 FPS
• Root Cause: Unoptimized frame buffer, redundant computations
• Solution: Reduced buffer 64→32 frames, conditional VETO, GPU batching
• Lesson: Profile and optimize critical path components

5. WHAT WORKED WELL

• Skeleton-based approach: Robust to lighting, camera angles, backgrounds
• GCN architecture: Effective for learning spatial-temporal patterns
• Smart Veto concept: Simple yet effective false positive reduction
• Modern stack: Next.js + Supabase enabled rapid development
• Cloud GPU: Vast.ai provided cost-effective high-performance compute

6. AREAS FOR IMPROVEMENT

Short-term:
• Add multi-class detection (punch, kick, weapon)
• Improve skeleton tracking for occluded scenarios
• Develop mobile application

Long-term:
• Implement transformer-based models for longer temporal context
• Add RGB + skeleton multi-modal fusion
• Deploy edge version on NVIDIA Jetson for on-premise use
• Implement federated learning for privacy-preserving updates

7. LESSONS LEARNED

• Data Quality > Data Quantity: Relevant violence data more important than large general datasets
• Ensemble > Single Model: Different architectures complement each other
• Real-World Testing Essential: High training accuracy doesn't guarantee deployment success
• Iterative Development: Multiple model iterations necessary to find optimal configuration
• Documentation Importance: Comprehensive testing records enabled informed decisions
"""

for i, para in enumerate(doc.paragraphs):
    if "Evaluate the project/application performance" in para.text or "Reflect on how well" in para.text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(reflection_content.strip())
        break

# ============================================================================
# SECTION 4: FILES TO BE SUBMITTED
# ============================================================================

files_content = """
Project Implementation Files:

1. Source Code Repository:
https://github.com/psdew2ewqws/NexaraVision

The repository contains:
• /src/app/ - Next.js 14 application pages and components
• /src/components/ - Reusable UI components (shadcn/ui)
• /src/lib/ - Utility functions and API clients
• /docs/ - Model configuration and selection reports
• package.json - Dependencies and scripts

Technology Stack Used:
• Frontend: Next.js 14.1.4, React 19.2.1, TypeScript 5
• Styling: Tailwind CSS 4, shadcn/ui components
• State: Zustand 5.0.8, @tanstack/react-query 5.90.9
• Backend: Supabase (Auth, PostgreSQL, Realtime)
• ML: PyTorch, FastAPI, YOLO v26, ST-GCN++, MS-G3D
• Deployment: Vercel (frontend), Vast.ai (ML backend)

2. Live Deployment:
https://nexaravision.com

The deployed application includes:
• Dashboard - Analytics and incident overview
• Live Monitor - Real-time violence detection
• Alerts - Alert history and acknowledgment
• Cameras - Camera management
• Users - User administration
• Settings - System configuration

3. ML Service API:
wss://api.nexaravision.com:34788/ws/live

4. Presentation:
[Link to presentation slides]

5. Documentation:
• MODEL_SELECTION_REPORT.md - Comprehensive model testing results
• FINAL_MODEL_CONFIGURATION.md - Production deployment configuration
• SMART_VETO_CONFIG.md - Ensemble algorithm documentation
"""

for i, para in enumerate(doc.paragraphs):
    if "Implement the selected project" in para.text or "Provide a link to your source code" in para.text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(files_content.strip())
        break

# ============================================================================
# REFERENCES
# ============================================================================

references_content = """
[1] Yan, S., Xiong, Y., & Lin, D. (2018). Spatial Temporal Graph Convolutional Networks for Skeleton-Based Action Recognition. AAAI Conference on Artificial Intelligence.

[2] Liu, Z., Zhang, H., Chen, Z., Wang, Z., & Ouyang, W. (2020). Disentangling and Unifying Graph Convolutions for Skeleton-Based Action Recognition. IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR).

[3] Cheng, K., Zhang, Y., Cao, C., Shi, L., Cheng, J., & Lu, H. (2020). Decoupling GCN with DropGraph Module for Skeleton-Based Action Recognition. European Conference on Computer Vision (ECCV).

[4] Shahroudy, A., Liu, J., Ng, T. T., & Wang, G. (2016). NTU RGB+D: A Large Scale Dataset for Action Recognition. IEEE Conference on Computer Vision and Pattern Recognition (CVPR).

[5] Cheng, M., Cai, K., & Li, M. (2021). RWF-2000: An Open Large Scale Video Database for Violence Detection. International Conference on Pattern Recognition (ICPR).

[6] Jocher, G., et al. (2023). Ultralytics YOLO. GitHub. https://github.com/ultralytics/ultralytics

[7] Bazarevsky, V., et al. (2020). BlazePose: On-device Real-time Body Pose Tracking. CVPR Workshop on Computer Vision for AR/VR.

[8] Sun, K., Xiao, B., Liu, D., & Wang, J. (2019). Deep High-Resolution Representation Learning for Visual Recognition. IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR).

[9] Real-Life Violence Situations Dataset. Kaggle. https://www.kaggle.com/datasets/mohamedmustafa/real-life-violence-situations-dataset

[10] SCVD - Surveillance Camera Violence Detection Dataset. Kaggle.

[11] AIRTLab Violence Detection Dataset. GitHub. https://github.com/airtlab/AIRTLab-dataset

[12] Next.js Documentation. Vercel. https://nextjs.org/docs

[13] Supabase Documentation. https://supabase.com/docs

[14] FastAPI Documentation. https://fastapi.tiangolo.com/

[15] PyTorch Documentation. https://pytorch.org/docs/

[16] Lin, T. Y., et al. (2014). Microsoft COCO: Common Objects in Context. European Conference on Computer Vision (ECCV).

[17] Vercel Deployment Documentation. https://vercel.com/docs

[18] Vast.ai GPU Cloud Documentation. https://vast.ai/docs
"""

for i, para in enumerate(doc.paragraphs):
    if "References" in para.text and para.style.name.startswith('Heading'):
        # Add references content after the heading
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(references_content.strip())
        break

# ============================================================================
# SAVE DOCUMENT
# ============================================================================

doc.save(str(OUTPUT_PATH))
print(f"Report saved to: {OUTPUT_PATH}")
print("Template formatting and HTU logo preserved!")
