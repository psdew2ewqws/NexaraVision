#!/usr/bin/env python3
"""
NexaraVision Capstone Report - CORRECTED VERSION
Uses actual validation data and correct tech stack from GitHub
Preserves HTU logo from original template
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Paths
TEMPLATE_PATH = Path('/home/admin/Downloads/Capstone Project 2  Template - students (2).docx')
OUTPUT_DIR = Path('/home/admin/Downloads/Final report')
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / 'NexaraVision_Capstone_Report.docx'
CHARTS_DIR = Path('/home/admin/Desktop/NexaraVision/claudedocs/report_charts')

# Load original template (preserves HTU logo)
doc = Document(str(TEMPLATE_PATH))

# ============================================================================
# HELPER: Find and replace text
# ============================================================================

def find_and_replace(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

# Update title page
find_and_replace(doc, "Your Project name", "NexaraVision")
find_and_replace(doc, "Student name and number", "Student Name - Student Number")
find_and_replace(doc, "Dr. name", "Dr. Supervisor Name")

# ============================================================================
# CONTENT SECTIONS (Corrected data)
# ============================================================================

# Section 1: Application Description
brief_overview = """
NexaraVision is a real-time violence detection system using skeleton-based Graph Convolutional Networks (GCNs). The system extracts human pose data using YOLO v26 and classifies violent behavior using ST-GCN++ and MS-G3D models in a Smart Veto Ensemble configuration.

TECHNOLOGY STACK (from package.json):
• Framework: Next.js 16.1.4 with App Router
• Language: TypeScript 5.x
• UI Library: React 19.2.1
• Styling: Tailwind CSS 4.x
• State Management: Zustand 5.0.8
• Data Fetching: @tanstack/react-query 5.90.9
• Backend: Supabase SSR 0.8.0, Supabase JS 2.91.0
• Animation: Framer Motion 12.27.5
• Charts: Recharts 3.4.1
• Internationalization: next-intl 4.7.0
• ML Backend: FastAPI + PyTorch 2.x
• Pose Estimation: YOLO v26 (yolo26m-pose.pt)
• Violence Detection: ST-GCN++ (PRIMARY) + MS-G3D (VETO)
• Deployment: Vercel (frontend) + Vast.ai RTX 4090 (ML backend)

KEY FEATURES:
• Real-Time Detection: ~50ms latency, ~20 FPS
• Smart Veto Ensemble: 0.1% false positive rate
• Multi-Camera Support: Screen recording segmentation
• Instant Alerts: WhatsApp, Telegram, Discord
• Evidence Recording: 30-second pre-buffer
• Web Dashboard: Modern Next.js interface
• Cloud Deployment: Scalable GPU infrastructure

PERFORMANCE METRICS:
• Training Accuracy: 94.56% (ST-GCN++), 95.17% (MS-G3D)
• Validation Accuracy: 95.41% (567 test samples)
• Test Score: 4/4 videos correct, 0 false positives
• Detection Threshold: 90% (optimal from 15 tested)
"""

# Section 1.2: Diagrams
diagrams_content = """
SYSTEM ARCHITECTURE:
[See Figure 1: System Architecture Diagram]

The system follows a microservices architecture:
1. Client Layer: Next.js 16 web application on Vercel
2. Backend Services: Supabase (Auth, PostgreSQL, Realtime)
3. ML Service: FastAPI on Vast.ai GPU (RTX 4090)

VIOLENCE DETECTION PIPELINE:
[See Figure 2: Pose Pipeline Flow]

Step 1: Video Frame Capture (Browser MediaDevices API)
Step 2: WebSocket Transmission (JPEG ~50KB)
Step 3: YOLO v26 Pose Estimation (17 COCO keypoints)
Step 4: Skeleton Normalization (hip-centered, 32-frame buffer)
Step 5: GCN Classification (ST-GCN++ PRIMARY + MSG3D VETO)
Step 6: Smart Veto Decision (PRIMARY≥94% AND VETO≥85%)
Step 7: Alert Dispatch (WhatsApp, Telegram, Discord)

SMART VETO ENSEMBLE:
[See Figure 3: Smart Veto Flow]

PRIMARY Model: STGCNPP_Kaggle_NTU (94.56% accuracy, 6.9MB)
VETO Model: MSG3D_Kaggle_NTU (95.17% accuracy, 4.2MB)
Decision: VIOLENCE if PRIMARY≥94% AND VETO≥85%
Result: 0.1% false positive rate (1/853 frames)

GCN MODEL ARCHITECTURES:
[See Figure 4: GCN Architecture]

ST-GCN++ (6 blocks): 3→64→64→128→128→256→256 channels
- Single 9-frame temporal convolution
- No dropout, 6.9MB model size
- Better at rejecting false positives

MS-G3D (6 blocks): Multi-scale temporal (3,5,7 frames)
- 30% dropout, 4.2MB model size
- Higher sensitivity to violence patterns
"""

# Section 1.3: Implementation Stages
stages_content = """
STAGE 1: Research & Dataset Collection (Weeks 1-2)
• Literature review on violence detection methods
• Selected skeleton-based approach for robustness
• Collected 133GB training data from multiple sources:
  - Kaggle Real-Life Violence (~2,000 videos)
  - NTU RGB+D 60 (56,578 skeleton samples)
  - RWF-2000 (1,980 surveillance videos)
  - SCVD (3,632 videos)
  - Hockey Fight (143 clips)
  - AIRTLab (384 videos)

STAGE 2: Skeleton Extraction Pipeline (Week 3)
• Implemented YOLO v26 pose estimation
• Developed hip-centered normalization
• Created COCO-to-NTU format conversion
• Extracted skeletons (~324,000 files)

STAGE 3: Model Training & Evaluation (Weeks 4-5)
• Trained 8 base models (4 datasets × 2 architectures)
• Trained 2 combined models (Kaggle+NTU)
• Configuration: AdamW, LR=1e-4, 30-50 epochs
• Key finding: NTU-only models failed (domain mismatch)

STAGE 4: Smart Veto Ensemble Development (Week 6)
• Analyzed model complementarity
• Tested 10 threshold combinations
• Selected: STGCNPP@94% + MSG3D@85%
• Achieved 0.1% false positive rate

STAGE 5: Web Application Development (Weeks 7-8)
• Frontend: Next.js 16, TypeScript, Tailwind CSS 4
• UI Components: shadcn/ui, Framer Motion
• State: Zustand + React Query
• Backend: Supabase Auth, PostgreSQL, Realtime
• Pages: Dashboard, Live Monitor, Alerts, Cameras, Settings

STAGE 6: Deployment & Production (Weeks 9-10)
• ML Backend: Vast.ai RTX 4090 (136.59.129.136:34788)
• Frontend: Vercel with auto-deployment
• Domain: nexaravision.com
• Live testing: Webcam, screen share, YouTube
"""

# Section 2: Test Plan
test_plan_content = """
TESTING METHODOLOGY:

1. UNIT TESTING
• Pose estimation accuracy (YOLO v26)
• GCN forward pass validation
• WebSocket message handling
• Alert dispatcher verification

2. MODEL VALIDATION TESTING
[See Figure 5: Validation Confusion Matrices]

ST-GCN++ (Option 2 - NTU Format):
• Dataset: 4,531 balanced NTU-format samples
• Test Set: 567 samples
• Validation Accuracy: 95.41%
• Confusion Matrix:
  - True Negatives: 252 (94.4% precision)
  - False Positives: 15
  - False Negatives: 11
  - True Positives: 289 (96.3% recall)

MSG3D (Option 1 - COCO Format):
• Dataset: 13,488 balanced COCO-format samples
• Test Set: 1,686 samples
• Validation Accuracy: 79.83%
• Confusion Matrix:
  - True Negatives: 661
  - False Positives: 164
  - False Negatives: 176
  - True Positives: 685

3. TEST VIDEO EVALUATION
[See Figure 6: Model Comparison]

4 Test Videos:
| Video | Expected | STGCNPP Score | Result |
|-------|----------|---------------|--------|
| Violence.mp4 | Violence | 97.4% | PASS |
| Non-Viol.mp4 | Non-Violence | 7.6% | PASS |
| Jewelry.mp4 | Non-Violence | 83.6% | PASS |
| Cereal.mp4 | Non-Violence | 0.5% | PASS |

Score: 4/4 correct, 0 false positives

4. THRESHOLD OPTIMIZATION
[See Figure 7: Threshold Analysis]

Tested 15 thresholds (30%-97%):
• Thresholds < 85%: Jewelry video causes FP
• Thresholds ≥ 85%: 4/4 correct, 0 FP
• Selected: 90% (6.4% margin below Jewelry)

5. SMART VETO LIVE TESTING
• Frames tested: 853
• False Positives: 1 (0.1%)
• Vetoed (caught FPs): 9
• True Negatives: 843

6. PERFORMANCE METRICS
| Metric | Target | Achieved |
|--------|--------|----------|
| Latency | <100ms | ~50ms |
| FPS | >10 | ~20 |
| FP Rate | <5% | 0.1% |
| Accuracy | >90% | 95.41% |
"""

# Section 2.2: Performance Analysis
performance_content = """
IMPLEMENTATION RESULTS:

1. MODEL TRAINING RESULTS
[See Figure 8: Training Curves]

Base Models (from MODEL_SELECTION_REPORT.md):
| Model | Dataset | Train Acc | Test Videos |
|-------|---------|-----------|-------------|
| MSG3D_Kaggle | Kaggle | 91.67% | 3/4 |
| MSG3D_NTU120 | NTU only | 94.18% | 0/4 (excluded) |
| STGCNPP_Kaggle | Kaggle | 91.99% | 3/4 |
| STGCNPP_NTU120 | NTU only | 95.43% | 0/4 (excluded) |

Combined Models:
| Model | Datasets | Train Acc | Test Videos |
|-------|----------|-----------|-------------|
| MSG3D_Kag+NTU | Kaggle+NTU | 95.17% | 2/4 |
| STGCNPP_Kag+NTU | Kaggle+NTU | 94.56% | 4/4 (SELECTED) |

2. KEY FINDINGS

Domain Mismatch Issue:
• NTU-only models achieved 0% on real-world videos
• High training accuracy (95%) ≠ real-world performance
• Solution: Combined Kaggle (violence) + NTU (normal actions)

Model Selection:
• STGCNPP_Kaggle+NTU selected despite lower training acc (94.56% vs 95.17%)
• Reason: Only model achieving 4/4 test videos with 0 FP
• MSG3D too aggressive (90.7% on non-violence video)

3. LATENCY BREAKDOWN
| Stage | Time |
|-------|------|
| Frame Capture | ~5ms |
| Network RTT | ~20ms |
| YOLO Pose | ~15ms |
| GCN PRIMARY | ~10ms |
| GCN VETO | ~8ms |
| Total | ~50ms |

4. OPTIMIZATIONS
• Reduced frame buffer: 64→32 frames
• Conditional VETO: Only runs when PRIMARY≥50%
• Binary WebSocket: 33% smaller than Base64
"""

# Section 3: Reflection
reflection_content = """
PROJECT EVALUATION:

1. OBJECTIVES ACHIEVED
| Requirement | Target | Achieved | Status |
|-------------|--------|----------|--------|
| Real-time detection | <100ms | ~50ms | EXCEEDED |
| Violence accuracy | >90% | 95.41% | EXCEEDED |
| False positive rate | <5% | 0.1% | EXCEEDED |
| Multi-camera support | Yes | Yes | MET |
| Alert notifications | 3 channels | 3 channels | MET |

2. TECHNICAL ACHIEVEMENTS
• Smart Veto Ensemble: Novel dual-model approach
• Domain Mismatch Solution: Combined dataset training
• End-to-End System: Camera to alert in production
• Modern Stack: Next.js 16, React 19, Tailwind 4

3. CHALLENGES & SOLUTIONS

Challenge 1: Domain Mismatch
• Problem: NTU-only models (95% acc) = 0% real-world
• Solution: Combined Kaggle+NTU training
• Result: 95.41% validation, 4/4 test videos

Challenge 2: False Positives
• Problem: Jewelry video (83.6%) triggered alerts
• Solution: 90% threshold + Smart Veto ensemble
• Result: 0.1% FP rate

Challenge 3: Model Selection Paradox
• Problem: Highest accuracy model had worst FP rate
• Solution: Use aggressive model as VETO only
• Result: Best of both approaches

4. LESSONS LEARNED
• Data domain must match deployment domain
• Ensemble > Single model for robustness
• Real-world testing essential
• High training accuracy ≠ deployment success

5. FUTURE IMPROVEMENTS
• Multi-class detection (punch, kick, weapon)
• Edge deployment (NVIDIA Jetson)
• Mobile application
• Transformer architectures
"""

# Section 4: Files
files_content = """
PROJECT FILES:

1. Source Code Repository:
https://github.com/psdew2ewqws/NexaraVision

Technology Stack (from package.json):
• next: 16.1.4
• react: 19.2.1
• typescript: 5.x
• tailwindcss: 4.x
• zustand: 5.0.8
• @tanstack/react-query: 5.90.9
• @supabase/supabase-js: 2.91.0
• framer-motion: 12.27.5
• recharts: 3.4.1
• next-intl: 4.7.0

2. Live Deployment:
https://nexaravision.com

3. ML Service API:
wss://136.59.129.136:34788/ws/live

4. Documentation:
• MODEL_SELECTION_REPORT.md - Model testing results
• TRAINING_RESULTS_SUMMARY.md - Validation metrics
• DATASET_CATALOG.md - Dataset documentation
"""

# References
references_content = """
[1] Yan, S., et al. (2018). Spatial Temporal Graph Convolutional Networks. AAAI.

[2] Liu, Z., et al. (2020). Disentangling and Unifying Graph Convolutions. CVPR.

[3] Shahroudy, A., et al. (2016). NTU RGB+D: Large Scale Dataset. CVPR.

[4] Cheng, M., et al. (2021). RWF-2000: Violence Detection Database. ICPR.

[5] Jocher, G., et al. (2023). Ultralytics YOLO. GitHub.

[6] Real-Life Violence Situations Dataset. Kaggle.

[7] Next.js Documentation. Vercel. https://nextjs.org/docs

[8] Supabase Documentation. https://supabase.com/docs

[9] FastAPI Documentation. https://fastapi.tiangolo.com/

[10] PyTorch Documentation. https://pytorch.org/docs/
"""

# ============================================================================
# UPDATE DOCUMENT CONTENT
# ============================================================================

# Find and update content sections
for i, para in enumerate(doc.paragraphs):
    text = para.text.lower()

    # Brief overview section
    if "brief overview about the application" in text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(brief_overview.strip())

    # Diagrams section
    elif "produce detailed diagrams" in text or "create visual representations" in text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(diagrams_content.strip())

    # Implementation stages
    elif "describe the different stages" in text or "clearly explain each phase" in text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(stages_content.strip())

    # Test plan
    elif "explain the testing procedures" in text or "describe how you tested" in text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(test_plan_content.strip())

    # Performance analysis
    elif "analyze the implementation results" in text or "review how the project performed" in text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(performance_content.strip())

    # Reflection
    elif "evaluate the project/application performance" in text or "reflect on how well" in text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(reflection_content.strip())

    # Files section
    elif "implement the selected project" in text or "provide a link to your source code" in text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(files_content.strip())

# Update references
for i, para in enumerate(doc.paragraphs):
    if "references" in para.text.lower() and para.style.name.startswith('Heading'):
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].clear()
            doc.paragraphs[i + 1].add_run(references_content.strip())
        break

# Save document
doc.save(str(OUTPUT_PATH))

print(f"{'='*60}")
print(f"CORRECTED Report saved to: {OUTPUT_PATH}")
print(f"HTU logo and template formatting preserved!")
print(f"{'='*60}")
print(f"\nCorrected data used:")
print(f"  - Tech Stack: Next.js 16.1.4, React 19.2.1, etc.")
print(f"  - Validation: 95.41% accuracy, 567 samples")
print(f"  - Test Videos: 4/4 correct, 0 FP")
print(f"{'='*60}")
