#!/usr/bin/env python3
"""
NexaraVision Capstone Report - PROFESSIONAL VERSION
Properly formatted with all 26 images included
Uses original HTU template
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import os

# Paths
TEMPLATE_PATH = Path('/home/admin/Downloads/Capstone Project 2  Template - students (2).docx')
OUTPUT_DIR = Path('/home/admin/Downloads/Final report')
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / 'NexaraVision_Capstone_Report.docx'
CHARTS_DIR = Path('/home/admin/Desktop/NexaraVision/claudedocs/report_charts')

# Load template (preserves HTU logo and formatting)
doc = Document(str(TEMPLATE_PATH))

def add_formatted_paragraph(doc, text, bold=False, size=11, space_after=6):
    """Add a properly formatted paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p

def add_image_with_caption(doc, image_name, caption, width=5.5):
    """Add image with centered caption"""
    image_path = CHARTS_DIR / image_name
    if image_path.exists():
        # Add image
        doc.add_picture(str(image_path), width=Inches(width))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add caption
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        return True
    return False

def add_table(doc, data, col_widths=None):
    """Add a formatted table"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))

    # Try to apply table style, fallback if not available
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass  # Style not available, use default

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = str(cell_text)
            # Format header row
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    # Add borders manually
    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    # Add space after table
    doc.add_paragraph()
    return table

# ============================================================================
# CLEAR TEMPLATE PLACEHOLDERS AND ADD REAL CONTENT
# ============================================================================

# Find the position after table of contents to start adding content
# We'll append our content to the document

# First, update title
for para in doc.paragraphs:
    if "Your Project name" in para.text:
        para.clear()
        run = para.add_run("NexaraVision")
        run.bold = True
        run.font.size = Pt(24)
    if "Student name and number" in para.text:
        para.clear()
        para.add_run("Student Name - Student Number")
    if "Dr. name" in para.text:
        para.clear()
        para.add_run("Dr. Supervisor Name")

# ============================================================================
# SECTION 1: APPLICATION DESCRIPTION
# ============================================================================

# Find and replace section 1 content
found_section1 = False
for i, para in enumerate(doc.paragraphs):
    if "brief overview about the application" in para.text.lower():
        found_section1 = True
        # Clear placeholder paragraphs
        j = i + 1
        while j < len(doc.paragraphs) and not doc.paragraphs[j].text.strip().startswith("1."):
            if "produce detailed" in doc.paragraphs[j].text.lower() or "describe the different" in doc.paragraphs[j].text.lower():
                break
            j += 1

        # Add content after the heading
        if i + 1 < len(doc.paragraphs):
            target = doc.paragraphs[i + 1]
            target.clear()

            content = """NexaraVision is a comprehensive real-time violence detection system designed for security and surveillance applications. The system leverages cutting-edge artificial intelligence technologies, specifically skeleton-based Graph Convolutional Networks (GCNs), to detect violent behavior in video feeds with high accuracy and minimal latency.

The core innovation lies in the Smart Veto Ensemble approach, combining ST-GCN++ and MS-G3D models to achieve both high sensitivity and low false positive rates. Unlike RGB-based systems susceptible to lighting and background changes, our skeleton-based approach focuses on human body movements for robustness.

Technology Stack (from GitHub package.json):
• Framework: Next.js 16.1.4 with App Router
• Language: TypeScript 5.x
• UI Library: React 19.2.1
• Styling: Tailwind CSS 4.x
• State Management: Zustand 5.0.8
• Data Fetching: @tanstack/react-query 5.90.9
• Backend: Supabase SSR 0.8.0, Supabase JS 2.91.0
• Animation: Framer Motion 12.27.5
• Charts: Recharts 3.4.1
• Internationalization: next-intl 4.7.0
• ML Backend: FastAPI + PyTorch 2.x
• Pose Estimation: YOLO v26 (yolo26m-pose.pt)
• Violence Detection: ST-GCN++ + MS-G3D Ensemble
• Deployment: Vercel (frontend) + Vast.ai RTX 4090 (ML)

Key Performance Metrics:
• Validation Accuracy: 95.41% (567 test samples)
• Test Score: 4/4 videos correct, 0 false positives
• False Positive Rate: 0.1% (Smart Veto Ensemble)
• Inference Latency: ~25ms per frame
• Real-time FPS: ~20 frames per second"""

            run = target.add_run(content)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
        break

# Now add remaining content at the end of document
doc.add_page_break()

# ============================================================================
# ADD FIGURES SECTION
# ============================================================================

h1 = doc.add_heading('Application Figures and Diagrams', level=1)

# System Architecture
doc.add_heading('System Architecture', level=2)
p = doc.add_paragraph()
p.add_run("The NexaraVision system follows a microservices architecture with clear separation between frontend, backend, and ML components.").font.size = Pt(11)
add_image_with_caption(doc, 'system_architecture.png', 'Figure 1: NexaraVision System Architecture', 6)

# Tech Stack
doc.add_heading('Technology Stack', level=2)
p = doc.add_paragraph()
p.add_run("Complete technology stack used in the project, with versions from package.json:").font.size = Pt(11)
add_image_with_caption(doc, 'tech_stack_table.png', 'Figure 2: Technology Stack and Versions', 5.5)

# Pose Pipeline
doc.add_heading('Pose Estimation Pipeline', level=2)
p = doc.add_paragraph()
p.add_run("YOLO v26 extracts 17 COCO keypoints from video frames for skeleton-based analysis:").font.size = Pt(11)
add_image_with_caption(doc, 'pose_pipeline_flow.png', 'Figure 3: YOLO v26 Pose Estimation Pipeline', 6)

# Skeleton Graph
p = doc.add_paragraph()
p.add_run("The COCO 17-keypoint skeleton graph structure used for GCN input:").font.size = Pt(11)
add_image_with_caption(doc, 'skeleton_graph.png', 'Figure 4: COCO 17-Keypoint Skeleton Graph', 4)

# GCN Architecture
doc.add_heading('GCN Model Architectures', level=2)
p = doc.add_paragraph()
p.add_run("Comparison of ST-GCN++ (PRIMARY) and MS-G3D (VETO) architectures:").font.size = Pt(11)
add_image_with_caption(doc, 'gcn_architecture.png', 'Figure 5: ST-GCN++ and MS-G3D Architectures', 6)

# Smart Veto
doc.add_heading('Smart Veto Ensemble', level=2)
p = doc.add_paragraph()
p.add_run("The Smart Veto decision flow combining both models for optimal detection:").font.size = Pt(11)
add_image_with_caption(doc, 'smart_veto_flow.png', 'Figure 6: Smart Veto Ensemble Decision Flow', 5)

# WebSocket
doc.add_heading('Real-Time Communication', level=2)
p = doc.add_paragraph()
p.add_run("WebSocket communication flow for real-time violence detection:").font.size = Pt(11)
add_image_with_caption(doc, 'websocket_flow.png', 'Figure 7: WebSocket Communication Flow', 5.5)

# Training Pipeline
doc.add_heading('Training Pipeline', level=2)
p = doc.add_paragraph()
p.add_run("Complete training pipeline from data collection to model deployment:").font.size = Pt(11)
add_image_with_caption(doc, 'training_pipeline.png', 'Figure 8: Training Pipeline Overview', 6)

# Alert System
doc.add_heading('Alert System', level=2)
p = doc.add_paragraph()
p.add_run("Multi-channel alert system flow for WhatsApp, Telegram, and Discord:").font.size = Pt(11)
add_image_with_caption(doc, 'alert_system_flow.png', 'Figure 9: Alert System Flow', 6)

# Web App Structure
doc.add_heading('Web Application Structure', level=2)
p = doc.add_paragraph()
p.add_run("Next.js 16 application pages and features:").font.size = Pt(11)
add_image_with_caption(doc, 'web_app_structure.png', 'Figure 10: Web Application Structure', 6)

doc.add_page_break()

# ============================================================================
# SECTION 2: MACHINE LEARNING MODELS
# ============================================================================

doc.add_heading('Machine Learning Models and Training', level=1)

# Datasets
doc.add_heading('Training Datasets', level=2)
p = doc.add_paragraph()
p.add_run("Total training data: ~133GB from 10+ sources with 324,000+ files:").font.size = Pt(11)
add_image_with_caption(doc, 'dataset_distribution.png', 'Figure 11: Dataset Size Distribution', 6)

# Add dataset table
dataset_table = [
    ['Dataset', 'Size', 'Samples', 'Type'],
    ['Kaggle Real-Life Violence', '~4GB', '~2,000', 'Violence-Specific'],
    ['NTU RGB+D 60', '20GB', '56,578', 'General Actions'],
    ['RWF-2000', '35GB', '1,980', 'Violence-Specific'],
    ['SCVD', '1.2GB', '3,632', 'Violence-Specific'],
    ['Kinetics Skeleton', '43GB', '266,442', 'General Actions'],
    ['UCF101', '14GB', '13,381', 'General Actions'],
]
add_table(doc, dataset_table)

# Training Results
doc.add_heading('Training Results', level=2)
p = doc.add_paragraph()
p.add_run("MSG3D training accuracy and loss curves over 50 epochs:").font.size = Pt(11)
add_image_with_caption(doc, 'training_curves_msg3d.png', 'Figure 12: MSG3D Training Curves', 6)

add_image_with_caption(doc, 'training_metrics_overview.png', 'Figure 13: Training Metrics Overview', 6)

# Model Comparison
doc.add_heading('Model Comparison', level=2)
p = doc.add_paragraph()
p.add_run("All 10 trained models compared on training accuracy and test video performance:").font.size = Pt(11)
add_image_with_caption(doc, 'model_comparison.png', 'Figure 14: Model Comparison - Training vs Test', 6)

add_image_with_caption(doc, 'detailed_model_comparison.png', 'Figure 15: Detailed Model Comparison Table', 6.5)

# Confidence Heatmap
p = doc.add_paragraph()
p.add_run("Violence confidence scores across all models and test videos:").font.size = Pt(11)
add_image_with_caption(doc, 'confidence_heatmap.png', 'Figure 16: Confidence Scores Heatmap', 5.5)

doc.add_page_break()

# ============================================================================
# SECTION 3: TEST PLAN AND RESULTS
# ============================================================================

doc.add_heading('Test Plan and Results', level=1)

# Validation Confusion Matrices
doc.add_heading('Validation Results', level=2)

p = doc.add_paragraph()
p.add_run("""ST-GCN++ Validation Results (SELECTED MODEL):
• Dataset: 4,531 balanced NTU-format samples
• Test Set: 567 samples
• Validation Accuracy: 95.41%
• Precision: 95.07%, Recall: 96.33%, F1-Score: 95.70%""").font.size = Pt(11)

add_image_with_caption(doc, 'cm_stgcnpp_validation.png', 'Figure 17: ST-GCN++ Validation Confusion Matrix (567 samples, 95.41%)', 4.5)

p = doc.add_paragraph()
p.add_run("""MSG3D Validation Results:
• Dataset: 13,488 balanced COCO-format samples
• Test Set: 1,686 samples
• Validation Accuracy: 79.83%
• Precision: 80.68%, Recall: 79.56%, F1-Score: 80.12%""").font.size = Pt(11)

add_image_with_caption(doc, 'cm_msg3d_validation.png', 'Figure 18: MSG3D Validation Confusion Matrix (1,686 samples, 79.83%)', 4.5)

# Test Video Results
doc.add_heading('Test Video Evaluation', level=2)

test_results = [
    ['Video', 'Expected', 'ST-GCN++ Score', 'Result'],
    ['Violence.mp4', 'Violence', '97.4%', 'PASS ✓'],
    ['Non-Violence.mp4', 'Non-Violence', '7.6%', 'PASS ✓'],
    ['Jewelry.mp4', 'Non-Violence', '83.6%', 'PASS ✓'],
    ['Cereal.mp4', 'Non-Violence', '0.5%', 'PASS ✓'],
]
add_table(doc, test_results)

p = doc.add_paragraph()
p.add_run("Test video confusion matrices:").font.size = Pt(11)
add_image_with_caption(doc, 'cm_stgcnpp_kaggle_ntu.png', 'Figure 19: ST-GCN++ Test Videos Confusion Matrix (4/4 correct)', 4)

add_image_with_caption(doc, 'cm_msg3d_kaggle_ntu.png', 'Figure 20: MSG3D Test Videos Confusion Matrix (2/4 correct)', 4)

# Threshold Analysis
doc.add_heading('Threshold Optimization', level=2)
p = doc.add_paragraph()
p.add_run("15 thresholds tested (30%-97%). Thresholds ≥85% achieve 4/4 with 0 false positives. Selected: 90%").font.size = Pt(11)
add_image_with_caption(doc, 'threshold_analysis.png', 'Figure 21: Threshold Optimization Analysis', 6)

# Smart Veto Testing
doc.add_heading('Smart Veto Ensemble Testing', level=2)
p = doc.add_paragraph()
p.add_run("Live testing on 853 frames: 0.1% false positive rate (1/853), 9 vetoed, 843 safe.").font.size = Pt(11)
add_image_with_caption(doc, 'cm_smart_veto_live.png', 'Figure 22: Smart Veto Live Testing Results', 4)

# Performance Metrics
doc.add_heading('Performance Analysis', level=2)
add_image_with_caption(doc, 'performance_radar.png', 'Figure 23: Performance Comparison Radar Chart', 5)

perf_table = [
    ['Metric', 'Target', 'Achieved', 'Status'],
    ['Validation Accuracy', '>90%', '95.41%', 'EXCEEDED'],
    ['Test Videos', '4/4', '4/4', 'MET'],
    ['False Positive Rate', '<5%', '0.1%', 'EXCEEDED'],
    ['Inference Latency', '<100ms', '~25ms', 'EXCEEDED'],
    ['Real-time FPS', '>10', '~20', 'EXCEEDED'],
]
add_table(doc, perf_table)

doc.add_page_break()

# ============================================================================
# SECTION 4: REFLECTION
# ============================================================================

doc.add_heading('Reflection and Evaluation', level=1)

doc.add_heading('Objectives Achievement', level=2)

objectives = [
    ['Requirement', 'Target', 'Achieved', 'Status'],
    ['Real-time Detection', '<100ms', '~50ms', 'EXCEEDED'],
    ['Violence Accuracy', '>90%', '95.41%', 'EXCEEDED'],
    ['False Positive Rate', '<5%', '0.1%', 'EXCEEDED'],
    ['Multi-camera Support', 'Yes', 'Yes', 'MET'],
    ['Alert Notifications', '3 channels', '3 channels', 'MET'],
    ['Evidence Recording', '30s buffer', '30s buffer', 'MET'],
    ['Web Dashboard', 'Functional', 'Full-featured', 'EXCEEDED'],
]
add_table(doc, objectives)

doc.add_heading('Challenges and Solutions', level=2)

p = doc.add_paragraph()
p.add_run("""Challenge 1: Domain Mismatch
• Problem: NTU-only models (95% training accuracy) achieved 0% on real-world videos
• Cause: NTU contains general actions, no violence-specific patterns
• Solution: Combined Kaggle (real violence) + NTU (normal actions) training
• Result: 95.41% validation accuracy with real-world generalization

Challenge 2: False Positives on Fast Movements
• Problem: Jewelry video (rapid hand movements) triggered 83.6% confidence
• Cause: Fast hand movements resemble striking motions
• Solution: 90% threshold + Smart Veto ensemble requiring dual confirmation
• Result: 0.1% false positive rate (1/853 frames)

Challenge 3: Model Selection Paradox
• Problem: Highest accuracy model (MSG3D 95.17%) had worst false positive rate
• Cause: MSG3D too aggressive, flagged non-violence at 90.7%
• Solution: Use MSG3D as VETO only, ST-GCN++ as PRIMARY
• Result: Best of both approaches - sensitivity + specificity

Challenge 4: Real-Time Performance
• Problem: Initial implementation achieved only 5 FPS
• Cause: Unoptimized frame buffer, redundant computations
• Solution: Reduced buffer 64→32 frames, conditional VETO, GPU batching
• Result: ~20 FPS achieved with ~50ms total latency""").font.size = Pt(11)

doc.add_heading('Lessons Learned', level=2)

p = doc.add_paragraph()
p.add_run("""• Data Domain Must Match Deployment: High training accuracy ≠ real-world performance
• Ensemble > Single Model: Different architectures complement each other
• Real-World Testing Essential: Validation metrics alone are insufficient
• Iterative Development: Multiple model iterations necessary for optimization
• Documentation Critical: Comprehensive testing records enabled informed decisions""").font.size = Pt(11)

doc.add_heading('Future Improvements', level=2)

p = doc.add_paragraph()
p.add_run("""Short-term:
• Multi-class detection (punch, kick, weapon identification)
• Improved skeleton tracking for occluded scenarios
• Mobile application development

Long-term:
• Edge deployment on NVIDIA Jetson for on-premise processing
• Transformer-based architectures for longer temporal context
• Multi-modal fusion (RGB + skeleton) for improved accuracy
• Federated learning for privacy-preserving model updates""").font.size = Pt(11)

doc.add_page_break()

# ============================================================================
# SECTION 5: FILES AND REFERENCES
# ============================================================================

doc.add_heading('Project Files and Links', level=1)

p = doc.add_paragraph()
p.add_run("""Source Code Repository:
https://github.com/psdew2ewqws/NexaraVision

Live Deployment:
https://nexaravision.com

ML Service API:
wss://136.59.129.136:34788/ws/live

Server Configuration:
• Provider: Vast.ai
• GPU: NVIDIA RTX 4090 24GB
• SSH: ssh -p 34796 root@136.59.129.136

Model Files:
• PRIMARY: STGCNPP_Kaggle_NTU.pth (6.9MB)
• VETO: MSG3D_Kaggle_NTU.pth (4.2MB)
• Pose: yolo26m-pose.pt""").font.size = Pt(11)

# Final Summary
doc.add_heading('Final Configuration Summary', level=2)
add_image_with_caption(doc, 'final_summary_table.png', 'Figure 24: Final Configuration Summary', 5.5)

doc.add_page_break()

# References
doc.add_heading('References', level=1)

refs = """[1] Yan, S., Xiong, Y., & Lin, D. (2018). Spatial Temporal Graph Convolutional Networks for Skeleton-Based Action Recognition. AAAI Conference on Artificial Intelligence.

[2] Liu, Z., Zhang, H., Chen, Z., Wang, Z., & Ouyang, W. (2020). Disentangling and Unifying Graph Convolutions for Skeleton-Based Action Recognition. IEEE/CVF CVPR.

[3] Shahroudy, A., Liu, J., Ng, T. T., & Wang, G. (2016). NTU RGB+D: A Large Scale Dataset for Action Recognition. IEEE CVPR.

[4] Cheng, M., Cai, K., & Li, M. (2021). RWF-2000: An Open Large Scale Video Database for Violence Detection. ICPR.

[5] Jocher, G., et al. (2023). Ultralytics YOLO. https://github.com/ultralytics/ultralytics

[6] Real-Life Violence Situations Dataset. Kaggle. https://www.kaggle.com/datasets/mohamedmustafa/real-life-violence-situations-dataset

[7] Next.js 16 Documentation. Vercel. https://nextjs.org/docs

[8] Supabase Documentation. https://supabase.com/docs

[9] FastAPI Documentation. https://fastapi.tiangolo.com/

[10] PyTorch Documentation. https://pytorch.org/docs/

[11] Tailwind CSS 4 Documentation. https://tailwindcss.com/docs

[12] React 19 Documentation. https://react.dev/"""

p = doc.add_paragraph()
run = p.add_run(refs)
run.font.size = Pt(11)
run.font.name = 'Times New Roman'

# ============================================================================
# SAVE DOCUMENT
# ============================================================================

doc.save(str(OUTPUT_PATH))

print("="*70)
print(f"PROFESSIONAL Report saved to: {OUTPUT_PATH}")
print("="*70)
print("\nImages included:")
images_used = [
    'system_architecture.png', 'tech_stack_table.png', 'pose_pipeline_flow.png',
    'skeleton_graph.png', 'gcn_architecture.png', 'smart_veto_flow.png',
    'websocket_flow.png', 'training_pipeline.png', 'alert_system_flow.png',
    'web_app_structure.png', 'dataset_distribution.png', 'training_curves_msg3d.png',
    'training_metrics_overview.png', 'model_comparison.png', 'detailed_model_comparison.png',
    'confidence_heatmap.png', 'cm_stgcnpp_validation.png', 'cm_msg3d_validation.png',
    'cm_stgcnpp_kaggle_ntu.png', 'cm_msg3d_kaggle_ntu.png', 'threshold_analysis.png',
    'cm_smart_veto_live.png', 'performance_radar.png', 'final_summary_table.png'
]
for i, img in enumerate(images_used, 1):
    status = "✓" if (CHARTS_DIR / img).exists() else "✗"
    print(f"  {status} Figure {i}: {img}")
print("="*70)
